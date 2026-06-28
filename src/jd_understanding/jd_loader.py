from pathlib import Path
from docx import Document


class JDLoader:
    """
    Job Description Loader

    Supports:
        • .docx
        • .txt

    Features:
        ✓ Table extraction
        ✓ Paragraph extraction
        ✓ Whitespace normalization
        ✓ Robust validation
    """

    SUPPORTED_EXTENSIONS = {".docx", ".txt"}

    def __init__(self, jd_path: str):

        self.jd_path = Path(jd_path)

        if not self.jd_path.exists():
            raise FileNotFoundError(
                f"Job description not found: {self.jd_path}"
            )

        if self.jd_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {self.jd_path.suffix}"
            )

    @staticmethod
    def _clean(text: str) -> str:

        return " ".join(text.split())

    def _load_docx(self) -> str:

        document = Document(self.jd_path)

        lines = []

        # -----------------------------
        # Paragraphs
        # -----------------------------
        for paragraph in document.paragraphs:

            text = self._clean(paragraph.text)

            if text:
                lines.append(text)

        # -----------------------------
        # Tables
        # -----------------------------
        for table in document.tables:

            for row in table.rows:

                row_text = " ".join(
                    self._clean(cell.text)
                    for cell in row.cells
                    if self._clean(cell.text)
                )

                if row_text:
                    lines.append(row_text)

        return "\n".join(lines)

    def _load_txt(self) -> str:

        with self.jd_path.open(
            "r",
            encoding="utf-8",
        ) as file:

            return self._clean(file.read())

    def load(self) -> str:

        if self.jd_path.suffix.lower() == ".docx":
            return self._load_docx()

        return self._load_txt()